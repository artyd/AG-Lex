"""Regenerate the Playwright test fixtures.

Run from the repo root:

    python e2e/fixtures/generate.py

Re-run only when you change the fixture contents — the produced .docx /
.xlsx are committed so e2e can run on a clean checkout without
python-docx/openpyxl installed.
"""
from pathlib import Path

from docx import Document
from openpyxl import Workbook


OUT = Path(__file__).resolve().parent


def write_contract():
    # Use "Стаття N" / "Розділ N" headings: the backend's split regex matches
    # those even when the docx→md converter escapes plain numbered headings
    # ("4.") as "4\." — escapes break the bare-number rule but not these.
    d = Document()
    d.add_heading("ДОГОВІР ПРО НАДАННЯ ПОСЛУГ", 0)
    d.add_paragraph("м. Київ, 01 квітня 2026 року")
    d.add_heading("Стаття 1 Предмет договору", level=2)
    d.add_paragraph(
        "Виконавець зобов'язується надати Замовнику послуги з юридичного супроводу."
    )
    d.add_heading("Стаття 4 Відповідальність сторін", level=2)
    d.add_paragraph(
        "У разі прострочення виконання зобов'язання Виконавець сплачує неустойку у "
        "розмірі 0,5% за кожен день прострочення від суми невиконаного зобов'язання."
    )
    d.add_heading("Стаття 6 Розірвання договору", level=2)
    d.add_paragraph(
        "Замовник має право розірвати договір у будь-який час шляхом письмового "
        "повідомлення Виконавця."
    )
    d.add_heading("Стаття 8 Форс-мажор", level=2)
    d.add_paragraph(
        "Сторони звільняються від відповідальності у разі настання обставин, що не "
        "залежать від їх волі."
    )
    d.add_heading("Стаття 10 Підписи сторін", level=2)
    d.add_paragraph("Замовник: _____________   Виконавець: _____________")
    d.save(OUT / "contract.docx")


def write_contract_pair():
    d = Document()
    d.add_heading("SUPPLY CONTRACT / ДОГОВІР ПОСТАЧАННЯ", 0)
    d.add_paragraph("Mumbai, India / Мумбаї, Індія — 01 April 2026")
    d.add_heading("2.3 Delivery basis / Базис постачання", level=2)
    d.add_paragraph("FCA Mumbai per Incoterms 2020. / FCA Мумбаї за Incoterms 2020.")
    d.add_heading("3.1 Price / Ціна", level=2)
    d.add_paragraph("USD 1 280 / metric ton")
    d.add_heading("4.1 Payment terms / Умови оплати", level=2)
    d.add_paragraph(
        "30% prepayment, 70% within 14 days of B/L. / 30% передоплата, 70% протягом "
        "14 днів з дати коносамента."
    )
    d.save(OUT / "contract-pair.docx")


def write_handover():
    wb = Workbook()
    ws = wb.active
    ws.title = "Таблиця 3"
    ws.append(["№", "Поле", "Значення"])
    ws.append(["1", "Постачальник", "KASYAP SWEETNERS PVT LTD"])
    ws.append(["2", "Продукт", "Sorbitol Solution 70% BP"])
    ws.append(["3", "Ціна", "USD 1 280/т"])
    ws.append(["4", "Валюта", "USD"])
    ws.append(["5", "Базис постачання", "CIF Odesa"])
    ws.append(["6", "Оплата", "100% протягом 30 днів"])
    wb.save(OUT / "handover.xlsx")


if __name__ == "__main__":
    write_contract()
    write_contract_pair()
    write_handover()
    print("Wrote:", *[p.name for p in OUT.glob("*") if p.suffix in {".docx", ".xlsx"}])
